"""
Document Processor: Ingests and chunks documents for RAG
Handles PDFs, text files, DOCX, and tabular data (XLSX, CSV)

⚠️  NOTICE: This is the basic document processor.
For enhanced document processing with superior PDF understanding, table extraction,
and multi-format support, use: rag/docling_processor.py

The Docling processor provides:
- Advanced PDF layout detection and reading order
- Structured table extraction with preserved formatting
- Support for PPTX, HTML, images, and OCR
- Better chunking quality for improved RAG retrieval

This basic processor is kept as a fallback and for backward compatibility.
"""

import os
import re
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document
from dataclasses import dataclass
import pandas as pd


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document"""
    content: str
    metadata: Dict[str, str]
    chunk_id: str


class DocumentProcessor:
    """
    Processes documents into chunks suitable for embedding and retrieval

    Dependencies:
    - PyPDF2: PDF text extraction
    - python-docx: DOCX text extraction
    - pandas: Tabular data processing (XLSX, CSV)
    - openpyxl: Excel file backend for pandas
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor
        
        Args:
            chunk_size: Target size of each chunk in characters
            chunk_overlap: Overlap between chunks to maintain context
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """
        Process all documents in a directory
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        directory = Path(directory_path)
        
        print(f"Processing documents from: {directory_path}")
        
        # Get all files first
        files = [f for f in directory.rglob('*') if f.is_file()]
        print(f"Found {len(files)} files")
        print()
        
        # Process all supported file types
        for i, file_path in enumerate(files, 1):
            print(f"[{i}/{len(files)}] Processing {file_path.name}...")
            
            if file_path.suffix.lower() == '.pdf':
                chunks.extend(self.process_pdf(str(file_path)))
            elif file_path.suffix.lower() in ['.txt', '.md']:
                chunks.extend(self.process_text_file(str(file_path)))
            elif file_path.suffix.lower() == '.docx':
                chunks.extend(self.process_docx(str(file_path)))
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                chunks.extend(self.process_excel(str(file_path)))
            elif file_path.suffix.lower() == '.csv':
                chunks.extend(self.process_csv(str(file_path)))
            else:
                print(f"  ⊘ Skipping unsupported file type: {file_path.suffix}")
        
        print()
        print(f"✅ Processed {len(chunks)} total chunks")
        return chunks
    
    def process_pdf(self, pdf_path: str) -> List[DocumentChunk]:
        """
        Extract and chunk text from PDF
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    full_text += f"\n[Page {page_num + 1}]\n{text}"
                
                # Create chunks with metadata
                text_chunks = self._chunk_text(full_text)
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append(DocumentChunk(
                        content=chunk_text,
                        metadata={
                            'source': os.path.basename(pdf_path),
                            'file_path': pdf_path,
                            'chunk_index': str(i),
                            'total_chunks': str(len(text_chunks)),
                            'file_type': 'pdf'
                        },
                        chunk_id=f"{os.path.basename(pdf_path)}_chunk_{i}"
                    ))
            
            print(f"  ✓ {os.path.basename(pdf_path)}: {len(chunks)} chunks")
            
        except Exception as e:
            print(f"  ✗ Error processing {pdf_path}: {e}")
        
        return chunks

    def process_docx(self, docx_path: str) -> List[DocumentChunk]:
        """
        Extract and chunk text from DOCX file

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of DocumentChunk objects
        """
        chunks = []

        try:
            doc = Document(docx_path)
            full_text = ""

            # Extract text from all paragraphs
            for para in doc.paragraphs:
                full_text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                full_text += "\n"

            # Create chunks with metadata
            text_chunks = self._chunk_text(full_text)

            for i, chunk_text in enumerate(text_chunks):
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    metadata={
                        'source': os.path.basename(docx_path),
                        'file_path': docx_path,
                        'chunk_index': str(i),
                        'total_chunks': str(len(text_chunks)),
                        'file_type': 'docx'
                    },
                    chunk_id=f"{os.path.basename(docx_path)}_chunk_{i}"
                ))

            print(f"  ✓ {os.path.basename(docx_path)}: {len(chunks)} chunks")

        except Exception as e:
            print(f"  ✗ Error processing {docx_path}: {e}")

        return chunks

    def process_text_file(self, file_path: str) -> List[DocumentChunk]:
        """
        Process plain text or markdown file
        
        Args:
            file_path: Path to text file
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                full_text = file.read()
                
                text_chunks = self._chunk_text(full_text)
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append(DocumentChunk(
                        content=chunk_text,
                        metadata={
                            'source': os.path.basename(file_path),
                            'file_path': file_path,
                            'chunk_index': str(i),
                            'total_chunks': str(len(text_chunks)),
                            'file_type': Path(file_path).suffix[1:]
                        },
                        chunk_id=f"{os.path.basename(file_path)}_chunk_{i}"
                    ))
            
            print(f"  ✓ {os.path.basename(file_path)}: {len(chunks)} chunks")
            
        except Exception as e:
            print(f"  ✗ Error processing {file_path}: {e}")
        
        return chunks
    
    def process_excel(self, excel_path: str) -> List[DocumentChunk]:
        """
        Extract and chunk data from Excel file
        
        Strategy for tables:
        - Convert each sheet to text format with headers
        - Preserve column relationships
        - Create chunks by row groups or semantic sections
        - Include sheet context in metadata
        
        Args:
            excel_path: Path to Excel file (.xlsx or .xls)
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(excel_path, engine='openpyxl')
            
            for sheet_name in excel_file.sheet_names:
                # Read sheet
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Skip empty sheets
                if df.empty:
                    continue
                
                # Convert dataframe to text chunks
                sheet_chunks = self._process_dataframe(
                    df,
                    source_file=os.path.basename(excel_path),
                    sheet_name=sheet_name,
                    file_path=excel_path
                )
                
                chunks.extend(sheet_chunks)
            
            print(f"  ✓ {os.path.basename(excel_path)}: {len(chunks)} chunks from {len(excel_file.sheet_names)} sheets")
            
        except Exception as e:
            print(f"  ✗ Error processing {excel_path}: {e}")
        
        return chunks
    
    def process_csv(self, csv_path: str) -> List[DocumentChunk]:
        """
        Extract and chunk data from CSV file
        
        Args:
            csv_path: Path to CSV file
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        try:
            # Read CSV
            df = pd.read_csv(csv_path)
            
            if df.empty:
                print(f"  ⊘ {os.path.basename(csv_path)}: Empty file")
                return chunks
            
            # Convert dataframe to text chunks
            chunks = self._process_dataframe(
                df,
                source_file=os.path.basename(csv_path),
                sheet_name='data',
                file_path=csv_path
            )
            
            print(f"  ✓ {os.path.basename(csv_path)}: {len(chunks)} chunks")
            
        except Exception as e:
            print(f"  ✗ Error processing {csv_path}: {e}")
        
        return chunks
    
    def _process_dataframe(
        self,
        df: pd.DataFrame,
        source_file: str,
        sheet_name: str,
        file_path: str
    ) -> List[DocumentChunk]:
        """
        Process a pandas DataFrame into text chunks
        
        Strategy:
        1. Convert table to structured text format
        2. Include column headers in each chunk
        3. Create chunks by row groups
        4. Preserve table structure and relationships
        
        Args:
            df: Pandas DataFrame
            source_file: Source filename
            sheet_name: Sheet/table name
            file_path: Full file path
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        # Get table metadata
        num_rows, num_cols = df.shape
        column_names = df.columns.tolist()
        
        # Strategy 1: For small tables (< 50 rows), create chunks by logical sections
        if num_rows <= 50:
            chunks.extend(
                self._chunk_small_table(df, source_file, sheet_name, file_path)
            )
        else:
            # Strategy 2: For large tables, create row-group chunks
            chunks.extend(
                self._chunk_large_table(df, source_file, sheet_name, file_path)
            )
        
        return chunks
    
    def _chunk_small_table(
        self,
        df: pd.DataFrame,
        source_file: str,
        sheet_name: str,
        file_path: str
    ) -> List[DocumentChunk]:
        """
        Process small table as a single or few chunks
        Preserves full table structure for better context
        """
        chunks = []
        
        # Create header section
        header_text = f"Table: {sheet_name}\n"
        header_text += f"Columns: {', '.join(df.columns.tolist())}\n"
        header_text += f"Rows: {len(df)}\n\n"
        
        # Convert table to markdown-style text
        table_text = self._dataframe_to_text(df)
        
        # Combine header and table
        full_text = header_text + table_text
        
        # Create chunks (may split if still too large)
        text_chunks = self._chunk_text(full_text)
        
        for i, chunk_text in enumerate(text_chunks):
            chunks.append(DocumentChunk(
                content=chunk_text,
                metadata={
                    'source': source_file,
                    'file_path': file_path,
                    'sheet_name': sheet_name,
                    'chunk_index': str(i),
                    'total_chunks': str(len(text_chunks)),
                    'file_type': 'excel',
                    'table_rows': str(len(df)),
                    'table_columns': str(len(df.columns)),
                    'columns': ', '.join(df.columns.tolist())
                },
                chunk_id=f"{source_file}_{sheet_name}_chunk_{i}"
            ))
        
        return chunks
    
    def _chunk_large_table(
        self,
        df: pd.DataFrame,
        source_file: str,
        sheet_name: str,
        file_path: str
    ) -> List[DocumentChunk]:
        """
        Process large table by creating chunks from row groups
        Each chunk includes column headers for context
        """
        chunks = []
        
        # Determine rows per chunk (aim for ~1000 chars)
        # Estimate: ~50 chars per row on average
        rows_per_chunk = max(10, self.chunk_size // 50)
        
        # Process in row groups
        for start_row in range(0, len(df), rows_per_chunk):
            end_row = min(start_row + rows_per_chunk, len(df))
            
            # Get row slice
            df_slice = df.iloc[start_row:end_row]
            
            # Create chunk with header
            chunk_text = f"Table: {sheet_name} (Rows {start_row+1}-{end_row})\n"
            chunk_text += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            chunk_text += self._dataframe_to_text(df_slice)
            
            chunks.append(DocumentChunk(
                content=chunk_text,
                metadata={
                    'source': source_file,
                    'file_path': file_path,
                    'sheet_name': sheet_name,
                    'chunk_index': str(len(chunks)),
                    'row_start': str(start_row + 1),
                    'row_end': str(end_row),
                    'total_rows': str(len(df)),
                    'file_type': 'excel',
                    'table_columns': str(len(df.columns)),
                    'columns': ', '.join(df.columns.tolist())
                },
                chunk_id=f"{source_file}_{sheet_name}_rows_{start_row+1}_{end_row}"
            ))
        
        return chunks
    
    def _dataframe_to_text(self, df: pd.DataFrame, max_width: int = 100) -> str:
        """
        Convert DataFrame to readable text format
        
        Uses pipe-separated format for clarity:
        Column1 | Column2 | Column3
        Value1  | Value2  | Value3
        
        Args:
            df: DataFrame to convert
            max_width: Maximum width per cell (truncate longer values)
            
        Returns:
            Formatted text representation
        """
        lines = []
        
        # Column headers
        headers = [str(col)[:max_width] for col in df.columns]
        lines.append(" | ".join(headers))
        lines.append("-" * (sum(len(h) for h in headers) + len(headers) * 3))
        
        # Data rows
        for idx, row in df.iterrows():
            # Convert each value to string, handle NaN
            values = []
            for val in row:
                if pd.isna(val):
                    values.append("")
                else:
                    val_str = str(val)[:max_width]
                    values.append(val_str)
            
            lines.append(" | ".join(values))
        
        return "\n".join(lines)
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        Optimized for handling very large documents efficiently
        
        Args:
            text: Full text to chunk
            
        Returns:
            List of text chunks
        """
        # Clean text first
        text = self._clean_text(text)
        
        chunks = []
        start = 0
        text_len = len(text)
        
        # Improved safety limit calculation for large documents
        # Allow 3x the theoretical maximum to handle edge cases gracefully
        expected_iterations = text_len // (self.chunk_size - self.chunk_overlap)
        max_iterations = max(expected_iterations * 3, 10000)  # At least 10k iterations for very large files
        iterations = 0
        
        while start < text_len:
            # Safety check to prevent infinite loops
            iterations += 1
            if iterations > max_iterations:
                completion_percent = (start / text_len) * 100
                print(f"⚠️ Warning: Hit iteration limit at position {start:,}/{text_len:,} ({completion_percent:.1f}% complete)")
                print(f"   Processed {len(chunks)} chunks so far. Consider increasing chunk_size for very large documents.")
                break
            
            # Find the end of the chunk
            end = start + self.chunk_size
            
            # If we're past the end, just take what's left
            if end >= text_len:
                end = text_len
            else:
                # Try to break at sentence boundary
                # Look for common sentence endings
                sentence_end = max(
                    text.rfind('. ', start, end),
                    text.rfind('.\n', start, end),
                    text.rfind('!\n', start, end),
                    text.rfind('?\n', start, end)
                )
                
                # Only use sentence boundary if it's found and reasonable
                # Ensure we have at least half the chunk size to avoid too-small chunks
                if sentence_end > start and sentence_end - start > self.chunk_size // 2:
                    end = sentence_end + 1
            
            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Move start position (with overlap)
            # Ensure we always make progress to avoid infinite loops
            if end > text_len - self.chunk_overlap:
                # Near the end, just finish
                break
            
            new_start = end - self.chunk_overlap
            
            # Critical: Always advance by at least some minimum amount
            min_advance = max(1, self.chunk_size // 4)  # At least 25% of chunk size
            if new_start <= start:
                new_start = start + min_advance
            elif new_start - start < min_advance:
                new_start = start + min_advance
            
            start = new_start
            
            # If we're at or past the end, we're done
            if start >= text_len:
                break
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        # Remove page numbers (common pattern)
        text = re.sub(r'\[Page \d+\]', '', text)
        
        # Remove excessive whitespace - optimized
        text = ' '.join(text.split())
        
        return text.strip()


# Example usage and testing
def main():
    """Test the document processor"""
    import sys
    
    if len(sys.argv) > 1:
        directory = sys.argv[1]
    else:
        directory = 'data/documents'
    
    processor = DocumentProcessor(chunk_size=1000, chunk_overlap=200)
    chunks = processor.process_directory(directory)
    
    print(f"\n✅ Total chunks created: {len(chunks)}")
    
    # Show sample chunk
    if chunks:
        print("\nSample chunk:")
        print(f"  ID: {chunks[0].chunk_id}")
        print(f"  Source: {chunks[0].metadata['source']}")
        print(f"  Content preview: {chunks[0].content[:200]}...")


if __name__ == "__main__":
    main()

