"""
Enhanced Markdown to Word Document Converter with AI Highlighting

Converts markdown documents to professional Microsoft Word format (.docx) using python-docx.
Supports:
- Headers (H1, H2, H3)
- Bold text
- Paragraphs
- Tables (pipe-delimited markdown tables)
- Horizontal rules
- Lists (bullet and numbered)
- AI-generated section highlighting (yellow)
- Track changes mode enabled

Author: DoD Contracting Automation System
Version: 1.1 (Simplified highlighting without text markers)

Dependencies:
- python-docx>=1.1.0: Word document creation and manipulation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # Add WD_COLOR_INDEX for highlighting
from docx.oxml.shared import OxmlElement, qn
import re
from typing import List, Tuple, Optional
from pathlib import Path


def is_table_row(line: str) -> bool:
    """
    Check if line is a markdown table row
    
    Args:
        line: Line to check
    
    Returns:
        True if line starts with | and contains at least one more |
    """
    stripped = line.strip()
    return stripped.startswith('|') and stripped.count('|') >= 2


def is_table_separator(line: str) -> bool:
    """
    Check if line is a markdown table separator (header separator)
    Example: |---|---|---|
    
    Args:
        line: Line to check
    
    Returns:
        True if line is a table separator
    """
    stripped = line.strip()
    if not stripped.startswith('|'):
        return False
    
    # Remove outer pipes and split
    inner = stripped.strip('|')
    cells = [cell.strip() for cell in inner.split('|')]
    
    # Check if all cells are made of dashes and colons
    for cell in cells:
        if not cell or not all(c in '-:' for c in cell):
            return False
    
    return True


def parse_table_row(line: str) -> List[str]:
    """
    Parse a markdown table row into cells
    
    Args:
        line: Table row line
    
    Returns:
        List of cell contents
    """
    # Remove outer pipes and split
    inner = line.strip().strip('|')
    cells = [cell.strip() for cell in inner.split('|')]
    return cells


def parse_markdown_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    """
    Parse a complete markdown table
    
    Args:
        lines: All markdown lines
        start_idx: Index where table starts
    
    Returns:
        Tuple of (table_data, end_idx)
        - table_data: 2D list of cell contents
        - end_idx: Index after last table row
    """
    table_data = []
    i = start_idx
    
    # Parse header row
    if i < len(lines) and is_table_row(lines[i]):
        header = parse_table_row(lines[i])
        table_data.append(header)
        i += 1
    
    # Skip separator row
    if i < len(lines) and is_table_separator(lines[i]):
        i += 1
    
    # Parse data rows
    while i < len(lines) and is_table_row(lines[i]):
        row = parse_table_row(lines[i])
        table_data.append(row)
        i += 1
    
    return table_data, i


def add_word_table(document: Document, table_data: List[List[str]], has_header: bool = True, highlight_ai: bool = False) -> None:
    """
    Add a table to the Word document
    
    Args:
        document: Word document object
        table_data: 2D list of cell contents
        has_header: Whether first row is a header
        highlight_ai: Whether to highlight as AI-generated (yellow)
    """
    if not table_data or not table_data[0]:
        return
    
    # Create table
    rows = len(table_data)
    cols = len(table_data[0])
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            
            # Remove markdown bold syntax and add text
            text = cell_data.replace('**', '')
            cell.text = text
            
            # Style header row
            if row_idx == 0 and has_header:
                # Make header bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        # Apply AI highlighting if requested
                        if highlight_ai:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # Shade header row
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E8E8E8')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Apply AI highlighting to data cells if requested
                if highlight_ai:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def apply_ai_highlighting(paragraph):
    """
    Apply yellow highlighting to all runs in a paragraph to mark AI-generated content
    
    Args:
        paragraph: Word paragraph object
    """
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def enable_track_changes(document: Document) -> None:
    """
    Enable track changes mode in the Word document
    
    Args:
        document: Word document object
    """
    # Access document settings to enable track revisions
    settings = document.settings
    settings_element = settings.element
    
    # Add track revisions setting
    track_revisions = OxmlElement('w:trackRevisions')
    settings_element.append(track_revisions)


def set_document_metadata(document: Document, title: str = None, program_name: str = None):
    """
    Set document metadata properties
    
    Args:
        document: Word document object
        title: Document title (from H1 header)
        program_name: Program name for subject field
    """
    core_properties = document.core_properties
    
    # Set metadata
    core_properties.author = "DoD Acquisition Automation System"
    core_properties.comments = ("Generated by AI - All content highlighted in YELLOW is AI-generated. "
                                "Review all sections before use. Track changes enabled for editing.")
    
    if title:
        core_properties.title = title
    
    if program_name:
        core_properties.subject = program_name
        core_properties.keywords = f"{program_name}, AI-Generated, Acquisition Document, Review Required"


def convert_markdown_bold_to_runs(paragraph, text: str, highlight_ai: bool = False) -> None:
    """
    Convert markdown **bold** syntax to Word bold runs
    
    Args:
        paragraph: Word paragraph object
        text: Text with markdown bold syntax
        highlight_ai: Whether to apply AI highlighting (yellow)
    """
    # Split text by ** markers
    parts = text.split('**')
    
    for i, part in enumerate(parts):
        if not part:
            continue
        
        run = paragraph.add_run(part)
        
        # Odd indices are bold (between ** markers)
        if i % 2 == 1:
            run.font.bold = True
        
        # Apply AI highlighting if requested
        if highlight_ai:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def parse_bullet_list(lines: List[str], start_idx: int) -> Tuple[List[str], int]:
    """
    Parse a bullet list
    
    Args:
        lines: All markdown lines
        start_idx: Index where list starts
    
    Returns:
        Tuple of (list_items, end_idx)
    """
    list_items = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('- ') or line.startswith('* '):
            list_items.append(line[2:])
            i += 1
        elif not line:
            i += 1
            break
        else:
            break
    
    return list_items, i


def parse_numbered_list(lines: List[str], start_idx: int) -> Tuple[List[str], int]:
    """
    Parse a numbered list
    
    Args:
        lines: All markdown lines
        start_idx: Index where list starts
    
    Returns:
        Tuple of (list_items, end_idx)
    """
    list_items = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        # Match patterns like "1. " or "1) "
        if re.match(r'^\d+[\.\)] ', line):
            # Extract text after number
            text = re.sub(r'^\d+[\.\)] ', '', line)
            list_items.append(text)
            i += 1
        elif not line:
            i += 1
            break
        else:
            break
    
    return list_items, i


def convert_markdown_to_docx(md_file: str, docx_file: str, program_name: str = None) -> None:
    """
    Convert markdown file to Word document with AI highlighting and track changes
    
    Args:
        md_file: Path to input markdown file
        docx_file: Path to output Word document file
        program_name: Optional program name for metadata
    
    Features:
        - Full markdown parsing (headers, tables, lists, bold text)
        - AI-generated section highlighting (yellow)
        - Track changes enabled
        - Professional formatting
        - Document metadata
    """
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    document = Document()
    
    # Track document title for metadata
    doc_title = None
    
    # Track AI-generated sections
    # Everything after an H2 header is considered AI-generated
    in_ai_section = False
    
    # Parse markdown line by line
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for table
        if is_table_row(line):
            # Parse complete table
            table_data, end_idx = parse_markdown_table(lines, i)
            
            # Add table to document with AI highlighting if in AI section
            if table_data:
                add_word_table(document, table_data, has_header=True, highlight_ai=in_ai_section)
                document.add_paragraph()  # Add spacing after table
            
            i = end_idx
            continue
        
        # H1 - Title
        if line.startswith('# '):
            text = line[2:].strip()
            doc_title = text
            
            paragraph = document.add_heading(text, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # H2 - Major Sections (start AI highlighting from here)
        elif line.startswith('## '):
            text = line[3:].strip()
            
            # Add heading
            heading = document.add_heading(text, level=2)
            
            # Mark that we're now in an AI-generated section
            in_ai_section = True
            
        # H3 - Subheadings
        elif line.startswith('### '):
            text = line[4:].strip()
            document.add_heading(text, level=3)
            
        # Horizontal rule
        elif line == '---':
            # Add a paragraph with bottom border to simulate HR
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            
        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            list_items, end_idx = parse_bullet_list(lines, i)
            
            for item in list_items:
                paragraph = document.add_paragraph(style='List Bullet')
                convert_markdown_bold_to_runs(paragraph, item, highlight_ai=in_ai_section)
            
            i = end_idx
            continue
        
        # Numbered list
        elif re.match(r'^\d+[\.\)] ', line):
            list_items, end_idx = parse_numbered_list(lines, i)
            
            for item in list_items:
                paragraph = document.add_paragraph(style='List Number')
                convert_markdown_bold_to_runs(paragraph, item, highlight_ai=in_ai_section)
            
            i = end_idx
            continue
        
        # Regular paragraph
        else:
            # Check if this is a bold-only line (metadata line)
            if line.startswith('**') and line.endswith('**'):
                text = line[2:-2]
                paragraph = document.add_paragraph()
                run = paragraph.add_run(text)
                run.font.bold = True
                
                # Apply AI highlighting if in AI section
                if in_ai_section:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # Continue reading multi-line paragraphs
                paragraph_lines = [line]
                i += 1
                
                while i < len(lines):
                    next_line = lines[i].strip()
                    # Stop at empty lines, headers, tables, or lists
                    if (not next_line or
                        next_line.startswith('#') or
                        next_line == '---' or
                        is_table_row(next_line) or
                        next_line.startswith('- ') or
                        next_line.startswith('* ') or
                        re.match(r'^\d+[\.\)] ', next_line)):
                        break
                    paragraph_lines.append(next_line)
                    i += 1
                
                # Join paragraph lines
                full_text = ' '.join(paragraph_lines)
                
                if full_text:
                    paragraph = document.add_paragraph()
                    convert_markdown_bold_to_runs(paragraph, full_text, highlight_ai=in_ai_section)
                
                continue
        
        i += 1
    
    # Enable track changes mode
    enable_track_changes(document)
    
    # Set document metadata with AI generation info
    set_document_metadata(document, title=doc_title, program_name=program_name)
    
    # Save document
    document.save(docx_file)
    
    print(f"✓ Word document created successfully: {docx_file}")
    print(f"  - Track changes enabled for editing")
    print(f"  - AI-generated sections highlighted in YELLOW")
    print(f"  - Professional formatting applied")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) >= 3:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
        program_name = sys.argv[3] if len(sys.argv) > 3 else None
    else:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx> [program_name]")
        print("\nExample:")
        print("  python convert_md_to_docx.py report.md report.docx 'ALMS Program'")
        sys.exit(1)
    
    try:
        convert_markdown_to_docx(md_file, docx_file, program_name)
    except FileNotFoundError:
        print(f"Error: File not found: {md_file}")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting markdown to Word: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
