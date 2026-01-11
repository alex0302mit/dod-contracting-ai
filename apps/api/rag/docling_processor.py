"""
Enhanced Document Processor using Docling

Replaces basic PyPDF2/python-docx processing with advanced parsing
that preserves layout, structure, and provides superior table extraction.

Key Features:
- Advanced PDF layout understanding (sections, headers, reading order)
- Structured table extraction with preserved formatting
- Multi-format support: PDF, DOCX, PPTX, XLSX, HTML, images
- OCR support for scanned documents
- Markdown export with preserved document structure

Dependencies:
- docling: Advanced document parsing and understanding
"""

import os
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass

# Patch torch.xpu before importing Docling to prevent 'module has no attribute xpu' error
# This is needed on macOS where Intel XPU support doesn't exist
try:
    import torch
    if not hasattr(torch, 'xpu'):
        # Create a mock xpu module that reports as unavailable
        class _MockXPU:
            @staticmethod
            def is_available():
                return False
            @staticmethod
            def device_count():
                return 0
        torch.xpu = _MockXPU()
except ImportError:
    pass  # torch not installed, Docling will handle this

# Import Docling components
# PdfFormatOption and PdfPipelineOptions allow CPU-only configuration
try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("⚠️  Warning: Docling not installed. Install with: pip install docling")


@dataclass
class DocumentChunk:
    """
    Represents a chunk of text from a document
    
    Maintains compatibility with existing RAG pipeline
    """
    content: str
    metadata: Dict[str, str]
    chunk_id: str


class DoclingProcessor:
    """
    Advanced document processor using Docling
    
    Features:
    - Advanced PDF layout understanding with reading order detection
    - Superior table extraction maintaining structure
    - Multi-format support (PDF, DOCX, PPTX, XLSX, HTML, images)
    - Better text extraction quality than basic libraries
    - Metadata extraction from document structure
    
    Dependencies:
    - docling: Advanced document parsing (pip install docling)
    
    Usage:
        processor = DoclingProcessor(chunk_size=1000, chunk_overlap=200)
        chunks = processor.process_document("path/to/document.pdf")
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize Docling processor
        
        Args:
            chunk_size: Target size for text chunks (characters)
            chunk_overlap: Overlap between chunks for context preservation
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize Docling converter once for efficiency
        if DOCLING_AVAILABLE:
            # Configure pipeline to use CPU only (avoids torch.xpu errors on macOS)
            # PdfPipelineOptions provides control over OCR and table extraction
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True  # Enable OCR for scanned documents
            pipeline_options.do_table_structure = True  # Enable table extraction
            
            # Create converter with PDF-specific options
            self.converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
        else:
            self.converter = None
            print("⚠️  Docling not available - will use fallback processor")
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """
        Process any supported document format using Docling
        
        Supported formats: PDF, DOCX, PPTX, XLSX, HTML, images, and more
        
        Args:
            file_path: Path to document
            
        Returns:
            List of DocumentChunk objects with enhanced metadata
        """
        file_name = Path(file_path).name
        print(f"  📄 Processing with Docling: {file_name}")
        
        # Check if Docling is available
        if not DOCLING_AVAILABLE or self.converter is None:
            print(f"    ⚠️  Docling not available, using fallback processor")
            return self._fallback_process(file_path)
        
        try:
            # Convert document - Docling handles format detection automatically
            result = self.converter.convert(file_path)
            doc = result.document
            
            # Extract markdown (preserves structure better than plain text)
            # Markdown maintains headers, lists, tables, and formatting
            markdown_text = doc.export_to_markdown()
            
            # Extract tables separately for better chunking and retrieval
            tables = self._extract_tables(doc)
            
            # Create chunks from markdown with table handling
            chunks = self._create_chunks(
                text=markdown_text,
                file_path=file_path,
                tables=tables
            )
            
            print(f"    ✓ Extracted {len(chunks)} chunks ({len(tables)} tables)")
            return chunks
            
        except Exception as e:
            print(f"    ⚠️  Error processing with Docling: {e}")
            print(f"    → Falling back to basic processing")
            # Fallback to basic processing if Docling fails
            return self._fallback_process(file_path)
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """
        Process all documents in a directory
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            All chunks from all processed documents
        """
        chunks = []
        directory = Path(directory_path)
        
        print(f"🔍 Scanning directory: {directory_path}")
        
        # Supported formats by Docling
        supported_extensions = {
            '.pdf', '.docx', '.pptx', '.xlsx', 
            '.html', '.htm', '.txt', '.md',
            '.png', '.jpg', '.jpeg', '.tiff', '.bmp'
        }
        
        # Find all supported files recursively
        files = [
            f for f in directory.rglob('*') 
            if f.is_file() and f.suffix.lower() in supported_extensions
        ]
        
        print(f"📚 Found {len(files)} supported documents")
        print()
        
        # Process each file
        for i, file_path in enumerate(files, 1):
            print(f"[{i}/{len(files)}]", end=" ")
            file_chunks = self.process_document(str(file_path))
            chunks.extend(file_chunks)
        
        print()
        print(f"✅ Total chunks extracted: {len(chunks)}")
        return chunks
    
    def _extract_tables(self, doc) -> List[str]:
        """
        Extract tables with structure preserved
        
        Docling provides superior table extraction compared to basic libraries.
        Tables are converted to markdown format to preserve structure.
        
        Args:
            doc: Docling document object
            
        Returns:
            List of table markdown strings
        """
        tables = []
        
        try:
            # Iterate through document elements to find tables
            for element in doc.body:
                # Check if element is a table
                if hasattr(element, 'label') and element.label == 'table':
                    # Export table as markdown to preserve structure
                    table_md = element.export_to_markdown()
                    if table_md and table_md.strip():
                        tables.append(table_md)
        except Exception as e:
            print(f"      Note: Table extraction warning: {e}")
        
        return tables
    
    def _create_chunks(
        self, 
        text: str, 
        file_path: str, 
        tables: List[str]
    ) -> List[DocumentChunk]:
        """
        Create overlapping chunks from text with smart sentence-based splitting
        
        Strategy:
        - Split by sentences for better semantic boundaries
        - Maintain chunk_size and chunk_overlap for consistency
        - Store tables as separate chunks with special metadata
        - Add processor metadata for tracking
        
        Args:
            text: Full document text (markdown format)
            file_path: Source file path
            tables: Extracted tables as markdown
            
        Returns:
            List of document chunks with metadata
        """
        chunks = []
        file_name = Path(file_path).name
        file_type = Path(file_path).suffix[1:].lower()
        
        # Split text into sentences for better chunking
        # This preserves semantic units better than character splitting
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        chunk_count = 0
        
        # Build chunks by adding sentences
        for sentence in sentences:
            # Check if adding this sentence exceeds chunk size
            if len(current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk += sentence + " "
            else:
                # Save current chunk if it has content
                if current_chunk.strip():
                    chunk_count += 1
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        metadata={
                            'source': file_name,
                            'file_path': file_path,
                            'chunk_index': str(chunk_count),
                            'file_type': file_type,
                            'content_type': 'text',
                            'processor': 'docling'  # Track which processor was used
                        },
                        chunk_id=f"{file_name}_docling_chunk_{chunk_count}"
                    ))
                
                # Start new chunk with overlap
                # Take last N characters for overlap to maintain context
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + sentence + " "
        
        # Add final chunk if there's remaining content
        if current_chunk.strip():
            chunk_count += 1
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                metadata={
                    'source': file_name,
                    'file_path': file_path,
                    'chunk_index': str(chunk_count),
                    'file_type': file_type,
                    'content_type': 'text',
                    'processor': 'docling'
                },
                chunk_id=f"{file_name}_docling_chunk_{chunk_count}"
            ))
        
        # Add tables as separate chunks for better retrieval
        # Tables often contain critical data (pricing, specifications, etc.)
        for i, table in enumerate(tables, 1):
            chunk_count += 1
            chunks.append(DocumentChunk(
                content=f"Table {i}:\n\n{table}",
                metadata={
                    'source': file_name,
                    'file_path': file_path,
                    'chunk_index': str(chunk_count),
                    'file_type': file_type,
                    'content_type': 'table',  # Mark as table for specialized retrieval
                    'table_number': str(i),
                    'processor': 'docling'
                },
                chunk_id=f"{file_name}_docling_table_{i}"
            ))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences using common sentence boundaries
        
        This is better than character-based splitting as it preserves
        semantic units and context.
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences
        """
        # Split on common sentence endings
        # Handle common abbreviations and edge cases
        import re
        
        # Replace newlines with spaces, but preserve paragraph breaks
        text = re.sub(r'\n\n+', ' [PARAGRAPH] ', text)
        text = re.sub(r'\n', ' ', text)
        
        # Split on sentence boundaries
        # Look for period, exclamation, or question mark followed by space and capital letter
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
        
        # Restore paragraph breaks
        sentences = [s.replace('[PARAGRAPH]', '\n\n') for s in sentences]
        
        return [s.strip() for s in sentences if s.strip()]
    
    def convert_to_html(self, file_path: str) -> Dict[str, any]:
        """
        Convert a document to HTML suitable for rich text editors

        This method exports the document as HTML rather than chunks,
        preserving formatting for use in Tiptap or similar editors.

        Args:
            file_path: Path to document (PDF, DOCX, etc.)

        Returns:
            Dict with 'html' content and 'warnings' list
        """
        file_name = Path(file_path).name
        print(f"  📄 Converting to HTML with Docling: {file_name}")
        warnings = []

        if not DOCLING_AVAILABLE or self.converter is None:
            # Fallback: basic text extraction wrapped in paragraphs
            print(f"    ⚠️  Docling not available, using fallback")
            warnings.append("Docling not available - using basic text extraction")
            return self._fallback_convert_to_html(file_path, warnings)

        try:
            # Convert document
            result = self.converter.convert(file_path)
            doc = result.document

            # Export to markdown first (preserves structure)
            markdown_text = doc.export_to_markdown()

            # Convert markdown to HTML
            html = self._markdown_to_html(markdown_text)

            print(f"    ✓ Converted to HTML ({len(html)} chars)")

            return {
                'html': html,
                'warnings': warnings
            }

        except Exception as e:
            print(f"    ⚠️  Error converting with Docling: {e}")
            warnings.append(f"Docling conversion error: {str(e)}")
            return self._fallback_convert_to_html(file_path, warnings)

    def _markdown_to_html(self, markdown_text: str) -> str:
        """
        Convert markdown to HTML suitable for Tiptap editor

        Args:
            markdown_text: Markdown formatted text

        Returns:
            HTML string
        """
        try:
            import markdown

            # Configure markdown processor with extensions for better conversion
            md = markdown.Markdown(extensions=[
                'tables',
                'fenced_code',
                'nl2br',  # Convert newlines to <br>
                'sane_lists',
            ])

            html = md.convert(markdown_text)

            # Ensure content is wrapped appropriately for Tiptap
            if not html.strip():
                return '<p></p>'

            # Enhance tables with styling for better rendering in Tiptap
            html = self._enhance_html_tables(html)

            # Detect bold paragraphs that should be headings
            html = self._detect_and_convert_headings(html)

            return html

        except ImportError:
            # Fallback: basic conversion without markdown library
            import re

            # Basic markdown to HTML conversion
            html = markdown_text

            # Headers
            html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
            html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
            html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

            # Bold and italic
            html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
            html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

            # Paragraphs
            paragraphs = html.split('\n\n')
            html = ''.join(
                f'<p>{p.strip()}</p>' if not p.strip().startswith('<') else p
                for p in paragraphs if p.strip()
            )

            # Apply enhancements even in fallback
            html = self._enhance_html_tables(html)
            html = self._detect_and_convert_headings(html)

            return html

    def _enhance_html_tables(self, html: str) -> str:
        """
        Add styling to HTML tables for better rendering in Tiptap editor

        Args:
            html: HTML content with tables

        Returns:
            HTML with enhanced table styling
        """
        import re

        # Add classes to table elements for consistent styling
        html = re.sub(r'<table>', '<table class="border-collapse w-full">', html)
        html = re.sub(r'<td>', '<td class="border border-slate-300 p-2">', html)
        html = re.sub(r'<td ', '<td class="border border-slate-300 p-2" ', html)
        html = re.sub(r'<th>', '<th class="border border-slate-300 p-2 bg-slate-100 font-semibold">', html)
        html = re.sub(r'<th ', '<th class="border border-slate-300 p-2 bg-slate-100 font-semibold" ', html)

        return html

    def _detect_and_convert_headings(self, html: str) -> str:
        """
        Detect bold-only paragraphs and convert to headings

        Heuristics:
        - Paragraph contains only <strong> or <b> tag
        - Text is short (< 100 chars)
        - Not ending with : (likely a label)
        - Not a single short word
        """
        import re

        # Pattern: <p><strong>Short text</strong></p>
        pattern = r'<p>\s*<(strong|b)>([^<]+)</\1>\s*</p>'

        def replace_heading(match):
            text = match.group(2).strip()

            # Skip if too long
            if len(text) > 100:
                return match.group(0)

            # Skip if ends with : (label)
            if text.endswith(':'):
                return match.group(0)

            # Skip single short words
            if ' ' not in text and len(text) < 15:
                return match.group(0)

            return f'<h2>{text}</h2>'

        return re.sub(pattern, replace_heading, html, flags=re.IGNORECASE)

    def _fallback_convert_to_html(self, file_path: str, warnings: List[str]) -> Dict[str, any]:
        """
        Fallback HTML conversion using basic text extraction

        Args:
            file_path: Path to file
            warnings: List to append warnings to

        Returns:
            Dict with 'html' and 'warnings'
        """
        try:
            suffix = Path(file_path).suffix.lower()
            text = ""

            if suffix == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            text += page.extract_text() + "\n\n"
                except Exception as e:
                    warnings.append(f"PDF extraction error: {str(e)}")

            elif suffix == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n\n"
                    # Extract tables
                    for table in doc.tables:
                        text += "\n"
                        for row in table.rows:
                            cells = [cell.text for cell in row.cells]
                            text += " | ".join(cells) + "\n"
                        text += "\n"
                except Exception as e:
                    warnings.append(f"DOCX extraction error: {str(e)}")

            elif suffix in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                warnings.append(f"Unsupported file type: {suffix}")
                return {'html': '<p>Unable to convert this file type.</p>', 'warnings': warnings}

            # Convert plain text to HTML paragraphs
            paragraphs = text.split('\n\n')
            html = ''.join(f'<p>{p.strip()}</p>' for p in paragraphs if p.strip())

            if not html:
                html = '<p>No text content could be extracted from this document.</p>'
                warnings.append("No text content extracted")

            return {'html': html, 'warnings': warnings}

        except Exception as e:
            warnings.append(f"Conversion failed: {str(e)}")
            return {'html': '<p>Failed to convert document.</p>', 'warnings': warnings}

    def _fallback_process(self, file_path: str) -> List[DocumentChunk]:
        """
        Fallback to basic processing if Docling fails or is unavailable
        
        Uses the existing basic document processor as a safety net.
        
        Args:
            file_path: Path to file
            
        Returns:
            Basic chunks using simple text extraction
        """
        # Import the basic processor
        try:
            from rag.document_processor import DocumentProcessor as BasicProcessor
            
            basic = BasicProcessor(self.chunk_size, self.chunk_overlap)
            
            # Route to appropriate basic method based on file type
            suffix = Path(file_path).suffix.lower()
            
            if suffix == '.pdf':
                return basic.process_pdf(file_path)
            elif suffix in ['.txt', '.md']:
                return basic.process_text_file(file_path)
            elif suffix == '.docx':
                return basic.process_docx(file_path)
            elif suffix in ['.xlsx', '.xls']:
                return basic.process_excel(file_path)
            elif suffix == '.csv':
                return basic.process_csv(file_path)
            else:
                print(f"      ✗ Unsupported file type: {suffix}")
                return []
                
        except Exception as e:
            print(f"      ✗ Fallback processing also failed: {e}")
            return []


# Example usage and testing
def main():
    """Test the Docling processor"""
    import sys
    
    if len(sys.argv) > 1:
        target = sys.argv[1]
    else:
        target = 'data/documents'
    
    # Initialize processor
    processor = DoclingProcessor(chunk_size=1000, chunk_overlap=200)
    
    # Check if target is file or directory
    target_path = Path(target)
    
    if target_path.is_file():
        # Process single file
        print(f"Processing single file: {target}")
        chunks = processor.process_document(str(target_path))
    elif target_path.is_dir():
        # Process directory
        print(f"Processing directory: {target}")
        chunks = processor.process_directory(target)
    else:
        print(f"❌ Path not found: {target}")
        return
    
    # Display results
    print(f"\n✅ Total chunks created: {len(chunks)}")
    
    # Show sample chunks
    if chunks:
        print("\nSample chunks:")
        for i, chunk in enumerate(chunks[:3], 1):
            print(f"\n{i}. ID: {chunk.chunk_id}")
            print(f"   Source: {chunk.metadata['source']}")
            print(f"   Type: {chunk.metadata.get('content_type', 'unknown')}")
            print(f"   Content preview: {chunk.content[:150]}...")


if __name__ == "__main__":
    main()

